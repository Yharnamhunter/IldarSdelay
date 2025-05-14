import os, io, zipfile, csv, json
from uuid import uuid4
from docx import Document
import time
from django.shortcuts import render, redirect
from .forms import RegisterForm, LoginForm
from .yandex_services import generate_text
from pathlib import Path
from django.utils.timezone import now
from .models import textGeneration, batchGeneration
from django.contrib.auth import login, logout, authenticate
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.http import FileResponse, HttpResponse, Http404, StreamingHttpResponse, HttpResponseBadRequest
from django.conf import settings
from django.urls import reverse
from django.utils.text import slugify

def register_view(request):
    if request.user.is_authenticated:
        return redirect('generate')

    if request.method == 'POST':
        form = RegisterForm(request.POST)
        if form.is_valid():
            user = form.save()

            raw_password = form.cleaned_data.get('password1')
            user = authenticate(
                username=user.username,
                password=raw_password
            )
            if user:
                login(request, user)
                messages.success(request, f"Добро пожаловать, {user.username}!")
                return redirect('generate')

            messages.error(request, "Пользователь создан, но не удалось войти автоматически.")
            return redirect('login')

        messages.error(request, "Пожалуйста, исправьте ошибки в форме.")
    else:
        form = RegisterForm()

    return render(request, 'generator/register.html', {
        'form': form
    })

def login_view(request):
    if request.user.is_authenticated:
        return redirect('generate')

    if request.method == 'POST':
        form = LoginForm(request, data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            messages.success(request, f'Добро пожаловать, {user.username}!')
            return redirect('generate')
        else:
            messages.error(request, 'Неправильный логин или пароль.')
    else:
        form = LoginForm()

    return render(request, 'generator/login.html', {'form': form})

@login_required
def logout_view(request):
    logout(request)
    messages.info(request, 'Вы успешно вышли из системы.')
    return redirect('login')

@login_required
def generate_view(request):
    """
    Одиночная генерация текста с проверкой длины (10–300 слов).
    """
    if request.method == 'POST':
        prompt = request.POST.get('prompt', '').strip()
        lang   = request.POST.get('lang', 'ru')
        try:
            length = int(request.POST.get('length', 120))
        except (TypeError, ValueError):
            length = 120
        if length < 10 or length > 300:
            messages.error(request, "Длина текста должна быть от 10 до 300 слов.")
            history = textGeneration.objects.filter(user=request.user).order_by('-created_at')
            return render(request, 'generator/generate.html', {
                'history': history,
                'prompt':  prompt,
                'lang':    lang,
                'length':  length,
            })

        # Генерация
        text = generate_text(prompt, lang, length)
        if not text:
            messages.error(request, "Не удалось сгенерировать текст. Попробуйте позже.")
            history = textGeneration.objects.filter(user=request.user).order_by('-created_at')
            return render(request, 'generator/generate.html', {
                'history': history,
                'prompt':  prompt,
                'lang':    lang,
                'length':  length,
            })

        # Сохранение и рендер результата
        textGeneration.objects.create(
            user     = request.user,
            prompt   = prompt,
            result   = text,
            language = lang
        )

        history = textGeneration.objects.filter(user=request.user).order_by('-created_at')
        return render(request, 'generator/generate.html', {
            'result': text,
            'prompt': prompt,
            'lang':   lang,
            'length': length,
            'history': history,
        })

    # GET
    history = textGeneration.objects.filter(user=request.user).order_by('-created_at')
    return render(request, 'generator/generate.html', {
        'history': history
    })


@login_required
def batch_generate_view(request):
    # Только GET-метод
    if request.method != 'GET':
        raise Http404()

    # История для боковой панели
    history = batchGeneration.objects.filter(
        user=request.user
    ).order_by('-created_at')[:10]

    # Если это просто заход на страницу — рендерим пустую форму
    if not request.GET:
        return render(request, 'generator/batch_generate.html', {
            'history': history,
            'initial': {}  # пустые initial, чтобы шаблон не падал
        })

    # Иначе — это уже "сабмит" формы, собираем параметры
    prompt   = request.GET.get('prompt', '').strip()
    try:
        count = int(request.GET.get('count', 5))
    except ValueError:
        count = 5

    lang    = request.GET.get('lang', 'ru')
    try:
        length = int(request.GET.get('length', 120))
    except ValueError:
        length = 120

    file_fmt = request.GET.get('file_format', 'txt').lower()

    # 2) Валидация — только prompt
    errors = []
    if not prompt:
        errors.append("Поле «Запрос» не может быть пустым.")

    if errors:
        for e in errors:
            messages.error(request, e)
        return render(request, 'generator/batch_generate.html', {
            'history': history,
            'initial': {
                'prompt': prompt,
                'count': count,
                'lang': lang,
                'length': length,
                'file_format': file_fmt
            }
        })

    # 3) Сохраняем в историю
    batchGeneration.objects.create(
        user=request.user,
        prompt=prompt,
        count=count,
        language=lang,
        file_format=file_fmt,
    )

    # 4) Генерируем тексты
    texts = []
    for i in range(count):
        txt = generate_text(prompt, language=lang)
        if not txt:
            messages.error(request, f"Ошибка генерации текста #{i+1}.")
            return render(request, 'generator/batch_generate.html', {
                'history': history
            })
        texts.append(txt)
        time.sleep(0.3)

    # 5) Формируем и отдаем файл
    timestamp = now().strftime("%Y%m%d_%H%M%S")
    filename  = f"batch_{timestamp}.{file_fmt}"

    if file_fmt == 'txt':
        buf = io.StringIO()
        for idx, t in enumerate(texts, start=1):
            buf.write(f"--- Text {idx} ---\n{t}\n\n")
        buf.seek(0)
        return FileResponse(
            buf,
            as_attachment=True,
            filename=filename,
            content_type='text/plain; charset=utf-8'
        )

    elif file_fmt == 'docx':
        doc = Document()
        for idx, t in enumerate(texts, start=1):
            doc.add_heading(f"Text {idx}", level=1)
            doc.add_paragraph(t)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return FileResponse(
            buf,
            as_attachment=True,
            filename=filename,
            content_type=(
                'application/vnd.openxmlformats-officedocument.'
                'wordprocessingml.document'
            )
        )

    elif file_fmt == 'rtf':
        buf = io.BytesIO()
        buf.write(b"{\\rtf1\\ansi\n")
        for idx, t in enumerate(texts, start=1):
            buf.write(f"\\b Text {idx} \\b0\\par\n".encode('utf-8'))
            safe = t.replace('\\','\\\\').replace('}','\\}')
            for line in safe.splitlines():
                buf.write((line + "\\par\n").encode('utf-8'))
            buf.write(b"\n")
        buf.write(b"}")
        buf.seek(0)
        return FileResponse(
            buf,
            as_attachment=True,
            filename=filename,
            content_type='application/rtf'
        )

    else:
        raise Http404("Неизвестный формат файла.")


@login_required
def batch_download_view(request, filename):

    file_path = os.path.join(settings.MEDIA_ROOT, 'batch', filename)

    if not os.path.exists(file_path):
        raise Http404("Файл не найден")

    return FileResponse(
        open(file_path, 'rb'),
        as_attachment=True,
        filename=filename
    )

@login_required
def batch_stream(request):
    prompt     = request.GET.get('prompt','').strip()
    lang       = request.GET.get('lang','ru')
    fmt        = request.GET.get('file_format','txt')
    try:
        count = max(1, min(20,int(request.GET.get('count',1))))
    except ValueError:
        count = 1
    try:
        max_tokens = int(request.GET.get('max_tokens',150))
    except ValueError:
        max_tokens = 150

    batchGeneration.objects.create(
        user=request.user,
        prompt=prompt,
        count=count,
        language=lang,
        file_format=fmt
    )

    base      = slugify(prompt) or 'text'
    uid       = uuid4().hex
    ext_map   = {
        'txt':  'txt',
        'docx': 'docx',
        'rtf':  'rtf',
        }
    ext       = ext_map.get(fmt,'txt')
    filename  = f"{base}_{uid}.{ext}"
    batch_dir = settings.MEDIA_ROOT / 'batch'
    batch_dir.mkdir(exist_ok=True)
    filepath  = batch_dir / filename

    def event_stream():
        texts = []
        for i in range(1, count+1):
            text = generate_text(prompt, lang, max_tokens)
            if not text:
                yield f"event: error\ndata: {json.dumps({'message':f'Сбой на шаге {i}'})}\n\n"
                return

            texts.append(text)
            pct = int(i/count*100)
            yield f"data: {json.dumps({'progress':pct})}\n\n"

        if fmt=='docx':
            doc = Document()
            for idx, t in enumerate(texts, start=1):
                doc.add_heading(f"Text {idx}", level=1)
                doc.add_paragraph(t)
            doc.save(filepath)
        elif fmt=='rtf':
            with open(filepath, 'wb') as f:
                f.write(b"{\\rtf1\\ansi\\ansicpg1251\\deff0\n")
                f.write(b"{\\fonttbl{\\f0\\fnil\\fcharset204 Times New Roman;}}\n")

                for idx, t in enumerate(texts, start=1):
                    header = f"\\pard\\sa200\\b Text {idx}\\b0\\par\n"
                    f.write(header.encode('cp1251'))
                    safe = (
                        t.replace('\\', r'\\').replace('{', r'\{').replace('}', r'\}')
                    )
                    for line in safe.splitlines():
                        f.write((line + "\\par\n").encode('cp1251'))
                f.write(b"\n")
                f.write(b"}")
        else:
            with open(filepath,'w',encoding='utf-8') as f:
                for idx,t in enumerate(texts,1):
                    f.write(f"--- Text {idx} ---\n{t}\n\n")

        download_url = reverse('batch_download', args=[filename])
        yield f"event: complete\ndata: {json.dumps({'download_url': download_url})}\n\n"

    return StreamingHttpResponse(event_stream(), content_type='text/event-stream')

@login_required
def clear_text_history(request):
    if request.method == "POST":
        textGeneration.objects.filter(user=request.user).delete()
        messages.success(request, "История одиночных запросов очищена.")
    return redirect('generate') 

@login_required
def clear_batch_history(request):
    if request.method == "POST":
        batchGeneration.objects.filter(user=request.user).delete()
        messages.success(request, "История пакетных запросов очищена.")
    return redirect('batch_generate')

def page_not_found_view(request, exception):
    return render(request, '404.html', status=404)